#! python
# USAGE: get the guestlists from 'guest.txt', and make new invitation letters just by one moment
# XI 2020 Arnold Cytrowski

import docx

with open('guests.txt') as source:
    names = source.readlines()
    document = docx.Document('document.docx')

    for name in names:
        name = name.strip()

        document.add_paragraph('\n\n\n\n\n\n\n\n')
        document.add_paragraph('I would like to invite', style='one')
        document.add_paragraph(name, style='name')
        document.add_paragraph('to a really funtastic party',
                               style='one')
        document.add_paragraph('in your\'s mama bedroom', style='one')
        document.add_paragraph("at 7 o'clock, today", style='one')

        document.add_page_break()

    document.save('invitations.docx')

    print('invitations are done, thank you')
    